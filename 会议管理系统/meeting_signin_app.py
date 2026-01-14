#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
会议纪要管理系统
使用PyQt6创建GUI界面，支持参会人员管理、签到跟踪、会议内容记录和Word文档导出
"""

import sys
import os
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QLineEdit, QTextEdit, QPushButton, QTableWidget,
                             QTableWidgetItem, QHeaderView, QComboBox, QDateTimeEdit,
                             QMessageBox, QFileDialog, QSplitter, QGroupBox, QFormLayout,
                             QCheckBox, QStatusBar)
from PyQt6.QtCore import Qt, QDateTime, pyqtSignal
from PyQt6.QtGui import QFont, QIcon, QAction
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json


class Attendee:
    """参会人员类"""
    def __init__(self, name="", department="", position="", contact="", status="未签到"):
        self.name = name
        self.department = department
        self.position = position
        self.contact = contact
        self.status = status
        self.signin_time = ""

    def to_dict(self):
        return {
            "name": self.name,
            "department": self.department,
            "position": self.position,
            "contact": self.contact,
            "status": self.status,
            "signin_time": self.signin_time
        }

    @classmethod
    def from_dict(cls, data):
        attendee = cls(
            data.get("name", ""),
            data.get("department", ""),
            data.get("position", ""),
            data.get("contact", ""),
            data.get("status", "未签到")
        )
        attendee.signin_time = data.get("signin_time", "")
        return attendee


class MeetingMinutesApp(QMainWindow):
    """会议纪要管理系统主窗口"""

    def __init__(self):
        super().__init__()
        self.attendees = []
        self.init_ui()
        self.load_data()

    def init_ui(self):
        """初始化用户界面"""
        self.setWindowTitle("会议纪要管理系统")
        self.setGeometry(100, 100, 1200, 800)

        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # 创建主布局
        main_layout = QHBoxLayout(central_widget)

        # 创建分割器
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)

        # 左侧面板 - 会议信息和人员管理
        left_panel = self.create_left_panel()
        splitter.addWidget(left_panel)

        # 右侧面板 - 签到表显示
        right_panel = self.create_right_panel()
        splitter.addWidget(right_panel)

        # 设置分割器比例
        splitter.setSizes([400, 800])

        # 创建菜单栏
        self.create_menu_bar()

        # 创建状态栏
        self.status_bar = self.statusBar()
        self.status_bar.showMessage("就绪")

        # 连接信号
        self.connect_signals()

    def create_left_panel(self):
        """创建左侧面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)

        # 会议信息组
        meeting_group = QGroupBox("会议信息")
        meeting_layout = QFormLayout(meeting_group)

        self.title_edit = QLineEdit("会议签到表")
        self.organizer_edit = QLineEdit()
        self.location_edit = QLineEdit()
        self.datetime_edit = QDateTimeEdit(QDateTime.currentDateTime())
        self.datetime_edit.setCalendarPopup(True)

        meeting_layout.addRow("会议标题:", self.title_edit)
        meeting_layout.addRow("主办单位:", self.organizer_edit)
        meeting_layout.addRow("会议地点:", self.location_edit)
        meeting_layout.addRow("会议时间:", self.datetime_edit)

        layout.addWidget(meeting_group)

        # 会议内容组
        content_group = QGroupBox("会议内容")
        content_layout = QVBoxLayout(content_group)

        # 会议主要内容
        content_layout.addWidget(QLabel("会议主要内容:"))
        self.main_content_edit = QTextEdit()
        self.main_content_edit.setPlaceholderText("请输入会议的主要议题和讨论内容...")
        self.main_content_edit.setMaximumHeight(100)
        content_layout.addWidget(self.main_content_edit)

        # 会议决定
        content_layout.addWidget(QLabel("会议决定:"))
        self.decisions_edit = QTextEdit()
        self.decisions_edit.setPlaceholderText("请输入会议做出的决定和决议...")
        self.decisions_edit.setMaximumHeight(100)
        content_layout.addWidget(self.decisions_edit)

        # 任务计划
        content_layout.addWidget(QLabel("任务计划:"))
        self.tasks_edit = QTextEdit()
        self.tasks_edit.setPlaceholderText("请输入后续任务安排和执行计划...")
        self.tasks_edit.setMaximumHeight(100)
        content_layout.addWidget(self.tasks_edit)

        layout.addWidget(content_group)

        # 添加人员组
        add_person_group = QGroupBox("添加参会人员")
        add_layout = QFormLayout(add_person_group)

        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText("请输入姓名")
        self.department_edit = QLineEdit()
        self.department_edit.setPlaceholderText("请输入部门")
        self.position_edit = QLineEdit()
        self.position_edit.setPlaceholderText("请输入职务")
        self.contact_edit = QLineEdit()
        self.contact_edit.setPlaceholderText("请输入联系方式")

        add_layout.addRow("姓名:", self.name_edit)
        add_layout.addRow("部门:", self.department_edit)
        add_layout.addRow("职务:", self.position_edit)
        add_layout.addRow("联系方式:", self.contact_edit)

        # 添加按钮
        self.add_button = QPushButton("添加人员")
        self.add_button.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; padding: 8px; }")
        add_layout.addRow(self.add_button)

        layout.addWidget(add_person_group)

        # 操作按钮组
        button_group = QGroupBox("操作")
        button_layout = QVBoxLayout(button_group)

        self.clear_button = QPushButton("清空列表")
        self.export_button = QPushButton("导出Word")
        self.export_button.setStyleSheet("QPushButton { background-color: #2196F3; color: white; padding: 10px; font-weight: bold; }")

        button_layout.addWidget(self.clear_button)
        button_layout.addWidget(self.export_button)

        layout.addWidget(button_group)

        # 添加伸缩空间
        layout.addStretch()

        return panel

    def create_right_panel(self):
        """创建右侧面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)

        # 签到表标题
        title_label = QLabel("参会人员签到表")
        title_label.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)

        # 添加会议纪要预览按钮
        self.preview_button = QPushButton("预览会议纪要")
        self.preview_button.setStyleSheet("QPushButton { background-color: #FF9800; color: white; padding: 8px; }")
        layout.addWidget(self.preview_button)

        # 创建表格
        self.table = QTableWidget()
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["姓名", "部门", "职务", "联系方式", "签到状态", "签到时间"])

        # 设置表格属性
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)  # 姓名列
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.Fixed)  # 签到状态列
        self.table.setColumnWidth(4, 100)
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.Fixed)  # 签到时间列
        self.table.setColumnWidth(5, 150)

        layout.addWidget(self.table)

        # 统计信息
        stats_group = QGroupBox("统计信息")
        stats_layout = QHBoxLayout(stats_group)

        self.total_label = QLabel("总人数: 0")
        self.signed_label = QLabel("已签到: 0")
        self.unsigned_label = QLabel("未签到: 0")

        stats_layout.addWidget(self.total_label)
        stats_layout.addWidget(self.signed_label)
        stats_layout.addWidget(self.unsigned_label)
        stats_layout.addStretch()

        layout.addWidget(stats_group)

        return panel

    def create_menu_bar(self):
        """创建菜单栏"""
        menubar = self.menuBar()

        # 文件菜单
        file_menu = menubar.addMenu("文件")

        save_action = QAction("保存", self)
        save_action.triggered.connect(self.save_data)
        file_menu.addAction(save_action)

        load_action = QAction("加载", self)
        load_action.triggered.connect(self.load_data_from_file)
        file_menu.addAction(load_action)

        file_menu.addSeparator()

        import_attendees_action = QAction("导入参会人员", self)
        import_attendees_action.triggered.connect(self.import_attendees)
        file_menu.addAction(import_attendees_action)

        export_attendees_action = QAction("导出参会人员", self)
        export_attendees_action.triggered.connect(self.export_attendees)
        file_menu.addAction(export_attendees_action)

        file_menu.addSeparator()

        exit_action = QAction("退出", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # 帮助菜单
        help_menu = menubar.addMenu("帮助")

        about_action = QAction("关于", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)

    def connect_signals(self):
        """连接信号"""
        self.add_button.clicked.connect(self.add_attendee)
        self.clear_button.clicked.connect(self.clear_attendees)
        self.export_button.clicked.connect(self.export_to_word)
        self.preview_button.clicked.connect(self.preview_minutes)
        self.table.cellDoubleClicked.connect(self.toggle_signin_status)

    def add_attendee(self):
        """添加参会人员"""
        name = self.name_edit.text().strip()
        if not name:
            QMessageBox.warning(self, "警告", "请输入姓名！")
            return

        department = self.department_edit.text().strip()
        position = self.position_edit.text().strip()
        contact = self.contact_edit.text().strip()

        attendee = Attendee(name, department, position, contact)
        self.attendees.append(attendee)

        self.update_table()
        self.clear_input_fields()

        self.status_bar.showMessage(f"已添加参会人员: {name}")

    def clear_input_fields(self):
        """清空输入字段"""
        self.name_edit.clear()
        self.department_edit.clear()
        self.position_edit.clear()
        self.contact_edit.clear()

    def update_table(self):
        """更新表格显示"""
        self.table.setRowCount(len(self.attendees))

        for row, attendee in enumerate(self.attendees):
            # 姓名
            name_item = QTableWidgetItem(attendee.name)
            self.table.setItem(row, 0, name_item)

            # 部门
            dept_item = QTableWidgetItem(attendee.department)
            self.table.setItem(row, 1, dept_item)

            # 职务
            pos_item = QTableWidgetItem(attendee.position)
            self.table.setItem(row, 2, pos_item)

            # 联系方式
            contact_item = QTableWidgetItem(attendee.contact)
            self.table.setItem(row, 3, contact_item)

            # 签到状态
            status_combo = QComboBox()
            status_combo.addItems(["未签到", "已签到", "请假"])
            status_combo.setCurrentText(attendee.status)
            status_combo.currentTextChanged.connect(lambda text, r=row: self.update_status(r, text))
            self.table.setCellWidget(row, 4, status_combo)

            # 签到时间
            time_item = QTableWidgetItem(attendee.signin_time)
            self.table.setItem(row, 5, time_item)

        self.update_statistics()

    def update_status(self, row, status):
        """更新签到状态"""
        if 0 <= row < len(self.attendees):
            old_status = self.attendees[row].status
            self.attendees[row].status = status

            if status == "已签到" and not self.attendees[row].signin_time:
                self.attendees[row].signin_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            elif status != "已签到":
                self.attendees[row].signin_time = ""

            self.table.item(row, 5).setText(self.attendees[row].signin_time)
            self.update_statistics()

            if status != old_status:
                self.status_bar.showMessage(f"{self.attendees[row].name} 状态更新为: {status}")

    def toggle_signin_status(self, row, column):
        """双击切换签到状态"""
        if column == 4:  # 签到状态列
            combo = self.table.cellWidget(row, 4)
            current_status = combo.currentText()
            if current_status == "未签到":
                combo.setCurrentText("已签到")
            elif current_status == "已签到":
                combo.setCurrentText("请假")
            else:
                combo.setCurrentText("未签到")

    def update_statistics(self):
        """更新统计信息"""
        total = len(self.attendees)
        signed = sum(1 for a in self.attendees if a.status == "已签到")
        unsigned = total - signed

        self.total_label.setText(f"总人数: {total}")
        self.signed_label.setText(f"已签到: {signed}")
        self.unsigned_label.setText(f"未签到: {unsigned}")

    def clear_attendees(self):
        """清空参会人员列表"""
        reply = QMessageBox.question(self, "确认", "确定要清空所有参会人员吗？",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.attendees.clear()
            self.update_table()
            self.status_bar.showMessage("已清空参会人员列表")

    def preview_minutes(self):
        """预览会议纪要"""
        preview_text = self.generate_preview_text()

        # 创建预览对话框
        from PyQt6.QtWidgets import QDialog, QTextEdit, QVBoxLayout, QHBoxLayout, QPushButton

        dialog = QDialog(self)
        dialog.setWindowTitle("会议纪要预览")
        dialog.setGeometry(200, 200, 800, 600)

        layout = QVBoxLayout(dialog)

        text_edit = QTextEdit()
        text_edit.setPlainText(preview_text)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)

        button_layout = QHBoxLayout()
        export_btn = QPushButton("导出Word")
        export_btn.clicked.connect(lambda: self.export_and_close(dialog))
        cancel_btn = QPushButton("关闭")
        cancel_btn.clicked.connect(dialog.close)

        button_layout.addStretch()
        button_layout.addWidget(export_btn)
        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)

        dialog.exec()

    def export_and_close(self, dialog):
        """导出并关闭预览对话框"""
        dialog.close()
        self.export_to_word()

    def generate_preview_text(self):
        """生成预览文本"""
        preview = []

        # 标题
        preview.append("=" * 50)
        preview.append(self.title_edit.text() or "会议纪要")
        preview.append("=" * 50)
        preview.append("")

        # 基本信息
        preview.append("会议基本信息:")
        preview.append(f"主办单位: {self.organizer_edit.text()}")
        preview.append(f"会议地点: {self.location_edit.text()}")
        preview.append(f"会议时间: {self.datetime_edit.dateTime().toString('yyyy-MM-dd hh:mm')}")
        preview.append("")

        # 参会统计
        total = len(self.attendees)
        signed = sum(1 for a in self.attendees if a.status == "已签到")
        unsigned = total - signed
        absent = sum(1 for a in self.attendees if a.status == "请假")

        preview.append("参会统计:")
        preview.append(f"总人数: {total}人，已签到: {signed}人，未签到: {unsigned}人，请假: {absent}人")
        preview.append("")

        # 会议主要内容
        main_content = self.main_content_edit.toPlainText().strip()
        if main_content:
            preview.append("一、会议主要内容")
            preview.append("-" * 30)
            preview.append(main_content)
            preview.append("")

        # 会议决定
        decisions = self.decisions_edit.toPlainText().strip()
        if decisions:
            preview.append("二、会议决定")
            preview.append("-" * 30)
            preview.append(decisions)
            preview.append("")

        # 任务计划
        tasks = self.tasks_edit.toPlainText().strip()
        if tasks:
            preview.append("三、任务计划")
            preview.append("-" * 30)
            preview.append(tasks)
            preview.append("")

        # 分页标记
        preview.append("")
        preview.append("=" * 50)
        preview.append("第2页：会议签到表")
        preview.append("=" * 50)
        preview.append("")

        # 签到表页面的会议信息
        preview.append("会议标题: " + (self.title_edit.text() or ""))
        preview.append("主办单位: " + (self.organizer_edit.text() or ""))
        preview.append("会议时间: " + self.datetime_edit.dateTime().toString("yyyy-MM-dd hh:mm"))
        preview.append("导出时间: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        preview.append("")

        # 参会统计
        preview.append("参会统计: " + f"总人数：{total}人，已签到：{signed}人，未签到：{unsigned}人，请假：{absent}人")
        preview.append("")

        # 参会人员名单
        if self.attendees:
            preview.append("参会人员签到表")
            preview.append("-" * 50)
            preview.append(f"{'序号':<4} {'姓名':<15} {'部门':<12} {'职务':<10} {'联系方式':<15} {'状态':<8}")
            preview.append("-" * 70)
            for i, attendee in enumerate(self.attendees, 1):
                status_mark = "✓" if attendee.status == "已签到" else ("✗" if attendee.status == "请假" else "○")
                time_info = f" ({attendee.signin_time})" if attendee.signin_time else ""
                preview.append(f"{i:<4} {attendee.name:<15} {attendee.department:<12} {attendee.position:<10} {attendee.contact:<15} {status_mark}{time_info}")

        return "\n".join(preview)

    def export_to_word(self):
        """导出为Word文档"""
        if not self.attendees:
            QMessageBox.warning(self, "警告", "没有参会人员数据，无法导出！")
            return

        filename, _ = QFileDialog.getSaveFileName(self, "保存Word文档",
                                                f"{self.title_edit.text() or '会议纪要'}.docx",
                                                "Word文档 (*.docx)")
        if not filename:
            return

        try:
            self.generate_word_document(filename)
            QMessageBox.information(self, "成功", f"已成功导出到: {filename}")
            self.status_bar.showMessage("Word文档导出成功")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")

    def generate_word_document(self, filename):
        """生成Word文档"""
        doc = Document()

        # 设置第一页标题为会议纪要
        title = doc.add_heading(self.title_edit.text() or "会议纪要", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加会议基本信息
        doc.add_paragraph("")  # 空行

        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'

        # 设置表格内容
        info_data = [
            ("主办单位", self.organizer_edit.text() or ""),
            ("会议地点", self.location_edit.text() or ""),
            ("会议时间", self.datetime_edit.dateTime().toString("yyyy-MM-dd hh:mm")),
            ("导出时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]

        for i, (label, value) in enumerate(info_data):
            cells = info_table.rows[i].cells
            cells[0].text = label + "："
            cells[1].text = value

        doc.add_paragraph("")  # 空行

        # 参会人员统计
        total = len(self.attendees)
        signed = sum(1 for a in self.attendees if a.status == "已签到")
        unsigned = total - signed
        absent = sum(1 for a in self.attendees if a.status == "请假")

        stats_para = doc.add_paragraph()
        stats_para.add_run("参会统计：").bold = True
        stats_para.add_run(f" 总人数：{total}人，已签到：{signed}人，未签到：{unsigned}人，请假：{absent}人")

        doc.add_paragraph("")  # 空行

        # 添加会议主要内容
        main_content = self.main_content_edit.toPlainText().strip()
        if main_content:
            content_heading = doc.add_heading("一、会议主要内容", level=2)
            content_para = doc.add_paragraph(main_content)
            content_para.paragraph_format.first_line_indent = Inches(0)  # 第一行无缩进

        # 添加会议决定
        decisions = self.decisions_edit.toPlainText().strip()
        if decisions:
            decisions_heading = doc.add_heading("二、会议决定", level=2)
            decisions_para = doc.add_paragraph(decisions)
            decisions_para.paragraph_format.first_line_indent = Inches(0)  # 第一行无缩进

        # 添加任务计划
        tasks = self.tasks_edit.toPlainText().strip()
        if tasks:
            tasks_heading = doc.add_heading("三、任务计划", level=2)
            tasks_para = doc.add_paragraph(tasks)
            tasks_para.paragraph_format.first_line_indent = Inches(0)  # 第一行无缩进

        doc.add_paragraph("")  # 空行

        # 添加分页符，进入新页面
        doc.add_page_break()

        # 创建签到表页面标题
        signin_title = doc.add_heading("会议签到表", level=1)
        signin_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 重新显示会议基本信息（在签到表页面）
        doc.add_paragraph("")  # 空行

        signin_info_table = doc.add_table(rows=4, cols=2)
        signin_info_table.style = 'Table Grid'

        # 设置表格内容（签到表页面的会议信息）
        signin_info_data = [
            ("会议标题", self.title_edit.text() or ""),
            ("主办单位", self.organizer_edit.text() or ""),
            ("会议时间", self.datetime_edit.dateTime().toString("yyyy-MM-dd hh:mm")),
            ("导出时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]

        for i, (label, value) in enumerate(signin_info_data):
            cells = signin_info_table.rows[i].cells
            cells[0].text = label + "："
            cells[1].text = value

        doc.add_paragraph("")  # 空行

        # 第二页的参会统计
        signin_stats_para = doc.add_paragraph()
        signin_stats_para.add_run("参会统计：").bold = True
        signin_stats_para.add_run(f"  总人数：{total}人，已签到：{signed}人，未签到：{unsigned}人，请假：{absent}人")

        doc.add_paragraph("")  # 空行

        signin_table = doc.add_table(rows=len(self.attendees) + 1, cols=6)
        signin_table.style = 'Table Grid'

        # 设置表头
        headers = ["序号", "姓名", "部门", "职务", "联系方式", "签到状态"]
        header_row = signin_table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # 设置表头样式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)

        # 添加参会人员数据
        for i, attendee in enumerate(self.attendees, 1):
            row = signin_table.rows[i]
            row.cells[0].text = str(i)
            row.cells[1].text = attendee.name
            row.cells[2].text = attendee.department
            row.cells[3].text = attendee.position
            row.cells[4].text = attendee.contact
            row.cells[5].text = attendee.status

        # 设置表格列宽
        for row in signin_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        # 添加页脚信息
        doc.add_paragraph("")  # 空行
        footer_para = doc.add_paragraph("注：本签到表由会议签到管理系统自动生成")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置整个文档的字体为宋体四号（14pt）
        self.set_document_font(doc)

        # 保存文档
        doc.save(filename)

    def set_document_font(self, doc):
        """设置整个文档的字体为宋体四号"""
        # 遍历所有段落设置字体
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(14)  # 四号字为14pt

        # 遍历所有表格设置字体
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(14)

    def save_data(self):
        """保存数据到文件"""
        data = {
            "meeting_info": {
                "title": self.title_edit.text(),
                "organizer": self.organizer_edit.text(),
                "location": self.location_edit.text(),
                "datetime": self.datetime_edit.dateTime().toString("yyyy-MM-dd hh:mm:ss"),
                "main_content": self.main_content_edit.toPlainText(),
                "decisions": self.decisions_edit.toPlainText(),
                "tasks": self.tasks_edit.toPlainText()
            },
            "attendees": [attendee.to_dict() for attendee in self.attendees]
        }

        filename, _ = QFileDialog.getSaveFileName(self, "保存数据", "meeting_data.json",
                                                "JSON文件 (*.json)")
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                self.status_bar.showMessage("数据保存成功")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"保存失败: {str(e)}")

    def load_data_from_file(self):
        """从文件加载数据"""
        filename, _ = QFileDialog.getOpenFileName(self, "加载数据", "", "JSON文件 (*.json)")
        if filename:
            try:
                with open(filename, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                # 加载会议信息
                meeting_info = data.get("meeting_info", {})
                self.title_edit.setText(meeting_info.get("title", ""))
                self.organizer_edit.setText(meeting_info.get("organizer", ""))
                self.location_edit.setText(meeting_info.get("location", ""))
                datetime_str = meeting_info.get("datetime", "")
                if datetime_str:
                    dt = QDateTime.fromString(datetime_str, "yyyy-MM-dd hh:mm:ss")
                    self.datetime_edit.setDateTime(dt)

                # 加载会议内容
                self.main_content_edit.setPlainText(meeting_info.get("main_content", ""))
                self.decisions_edit.setPlainText(meeting_info.get("decisions", ""))
                self.tasks_edit.setPlainText(meeting_info.get("tasks", ""))

                # 加载参会人员
                self.attendees = []
                for attendee_data in data.get("attendees", []):
                    attendee = Attendee.from_dict(attendee_data)
                    self.attendees.append(attendee)

                self.update_table()
                self.status_bar.showMessage("数据加载成功")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"加载失败: {str(e)}")

    def import_attendees(self):
        """导入参会人员"""
        filename, _ = QFileDialog.getOpenFileName(self, "导入参会人员", "", "CSV文件 (*.csv);;Excel文件 (*.xlsx);;文本文件 (*.txt)")
        if not filename:
            return

        try:
            imported_count = 0
            if filename.endswith('.csv'):
                imported_count = self.import_from_csv(filename)
            elif filename.endswith('.xlsx'):
                imported_count = self.import_from_excel(filename)
            elif filename.endswith('.txt'):
                imported_count = self.import_from_txt(filename)

            if imported_count > 0:
                self.update_table()
                QMessageBox.information(self, "成功", f"成功导入 {imported_count} 位参会人员")
                self.status_bar.showMessage(f"导入 {imported_count} 位参会人员成功")
            else:
                QMessageBox.warning(self, "警告", "没有导入任何参会人员数据")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导入失败: {str(e)}")

    def export_attendees(self):
        """导出参会人员"""
        if not self.attendees:
            QMessageBox.warning(self, "警告", "没有参会人员数据，无法导出！")
            return

        filename, file_type = QFileDialog.getSaveFileName(self, "导出参会人员", "参会人员.csv",
                                                        "CSV文件 (*.csv);;Excel文件 (*.xlsx);;文本文件 (*.txt)")
        if not filename:
            return

        try:
            if filename.endswith('.csv'):
                self.export_to_csv(filename)
            elif filename.endswith('.xlsx'):
                self.export_to_excel(filename)
            elif filename.endswith('.txt'):
                self.export_to_txt(filename)

            QMessageBox.information(self, "成功", f"参会人员已导出到: {filename}")
            self.status_bar.showMessage("参会人员导出成功")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")

    def import_from_csv(self, filename):
        """从CSV文件导入参会人员"""
        import csv
        count = 0
        with open(filename, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            for row in reader:
                name = row.get('姓名', '').strip()
                if name:  # 只导入有姓名的记录
                    attendee = Attendee(
                        name=name,
                        department=row.get('部门', '').strip(),
                        position=row.get('职务', '').strip(),
                        contact=row.get('联系方式', '').strip(),
                        status=row.get('签到状态', '未签到').strip()
                    )
                    self.attendees.append(attendee)
                    count += 1
        return count

    def import_from_excel(self, filename):
        """从Excel文件导入参会人员"""
        try:
            import pandas as pd
            df = pd.read_excel(filename)
            count = 0
            for _, row in df.iterrows():
                name = str(row.get('姓名', '')).strip()
                if name:  # 只导入有姓名的记录
                    attendee = Attendee(
                        name=name,
                        department=str(row.get('部门', '')).strip(),
                        position=str(row.get('职务', '')).strip(),
                        contact=str(row.get('联系方式', '')).strip(),
                        status=str(row.get('签到状态', '未签到')).strip()
                    )
                    self.attendees.append(attendee)
                    count += 1
            return count
        except ImportError:
            raise ImportError("需要安装pandas库才能导入Excel文件: pip install pandas openpyxl")

    def import_from_txt(self, filename):
        """从文本文件导入参会人员"""
        count = 0
        with open(filename, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):  # 跳过空行和注释行
                    parts = line.split(',')
                    if len(parts) >= 1:
                        name = parts[0].strip()
                        if name:
                            attendee = Attendee(
                                name=name,
                                department=parts[1].strip() if len(parts) > 1 else "",
                                position=parts[2].strip() if len(parts) > 2 else "",
                                contact=parts[3].strip() if len(parts) > 3 else "",
                                status=parts[4].strip() if len(parts) > 4 else "未签到"
                            )
                            self.attendees.append(attendee)
                            count += 1
        return count

    def export_to_csv(self, filename):
        """导出参会人员到CSV文件"""
        import csv
        with open(filename, 'w', newline='', encoding='utf-8-sig') as f:
            fieldnames = ['姓名', '部门', '职务', '联系方式', '签到状态', '签到时间']
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for attendee in self.attendees:
                writer.writerow({
                    '姓名': attendee.name,
                    '部门': attendee.department,
                    '职务': attendee.position,
                    '联系方式': attendee.contact,
                    '签到状态': attendee.status,
                    '签到时间': attendee.signin_time
                })

    def export_to_excel(self, filename):
        """导出参会人员到Excel文件"""
        try:
            import pandas as pd
            data = []
            for attendee in self.attendees:
                data.append({
                    '姓名': attendee.name,
                    '部门': attendee.department,
                    '职务': attendee.position,
                    '联系方式': attendee.contact,
                    '签到状态': attendee.status,
                    '签到时间': attendee.signin_time
                })
            df = pd.DataFrame(data)
            df.to_excel(filename, index=False)
        except ImportError:
            raise ImportError("需要安装pandas库才能导出Excel文件: pip install pandas openpyxl")

    def export_to_txt(self, filename):
        """导出参会人员到文本文件"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("# 参会人员列表\n")
            f.write("# 格式: 姓名,部门,职务,联系方式,签到状态\n")
            for attendee in self.attendees:
                line = f"{attendee.name},{attendee.department},{attendee.position},{attendee.contact},{attendee.status}"
                if attendee.signin_time:
                    line += f",{attendee.signin_time}"
                f.write(line + "\n")

    def load_data(self):
        """加载默认数据（如果存在）"""
        try:
            if os.path.exists("meeting_data.json"):
                with open("meeting_data.json", 'r', encoding='utf-8') as f:
                    data = json.load(f)

                # 加载会议信息
                meeting_info = data.get("meeting_info", {})
                self.title_edit.setText(meeting_info.get("title", "会议签到表"))
                self.organizer_edit.setText(meeting_info.get("organizer", ""))
                self.location_edit.setText(meeting_info.get("location", ""))
                datetime_str = meeting_info.get("datetime", "")
                if datetime_str:
                    dt = QDateTime.fromString(datetime_str, "yyyy-MM-dd hh:mm:ss")
                    self.datetime_edit.setDateTime(dt)

                # 加载会议内容
                self.main_content_edit.setPlainText(meeting_info.get("main_content", ""))
                self.decisions_edit.setPlainText(meeting_info.get("decisions", ""))
                self.tasks_edit.setPlainText(meeting_info.get("tasks", ""))

                # 加载参会人员
                self.attendees = []
                for attendee_data in data.get("attendees", []):
                    attendee = Attendee.from_dict(attendee_data)
                    self.attendees.append(attendee)

                self.update_table()
        except:
            pass  # 忽略加载错误

    def show_about(self):
        """显示关于对话框"""
        QMessageBox.about(self, "关于",
                         "会议纪要管理系统 v1.0\n\n"
                         "基于PyQt6开发的完整会议管理工具\n"
                         "支持人员管理、签到跟踪、会议内容记录和Word文档导出功能")


def main():
    """主函数"""
    app = QApplication(sys.argv)

    # 设置应用程序信息
    app.setApplicationName("会议纪要管理系统")
    app.setApplicationVersion("1.0")
    app.setOrganizationName("会议管理系统")

    # 设置样式
    app.setStyle("Fusion")

    # 创建主窗口
    window = MeetingMinutesApp()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()